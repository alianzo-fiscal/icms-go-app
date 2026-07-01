#!/usr/bin/env python3
"""
analisar_saidas_df.py
Analise de Saidas ICMS/DF (Distrito Federal) -- gera Excel (.xlsx) e Word (.docx)
Uso: python analisar_saidas_df.py <arquivo1.csv> [arquivo2.xls] ...
"""

from __future__ import annotations
import sys
import datetime
from pathlib import Path

import pandas as pd

PASTA_SAIDA = Path(__file__).parent
ALIQ_INT_DF   = 20.0          # RICMS/DF (Decreto 18.955/97) + Lei 7.326/2023 -- desde 21/01/2024
ALIQ_ANT_DF   = 18.0          # aliquota anterior
DATA_MUDANCA  = datetime.date(2024, 1, 21)

MESES_PT = {
    1: 'Janeiro', 2: 'Fevereiro', 3: 'Marco', 4: 'Abril',
    5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
    9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
}

COLS_DIV = [
    'CODFILIAL', 'NUMNOTA', 'DATA', 'SITTRIBUT', 'CODFISCAL',
    'ORIGMERCTRIB', 'CODPROD', 'DESCRICAO', 'NCM',
    'VLITEM', 'VLBASEICMS', 'PERCICMS', 'VLICMS',
    'ICMS Correto (R$)', 'Dif. ICMS (R$)', 'Critica', '_arquivo'
]


# -- Leitura -------------------------------------------------------------------
def carregar_dados(caminhos):
    dfs = []
    for caminho in caminhos:
        p = Path(caminho)
        ext = p.suffix.lower()
        try:
            if ext == '.csv':
                d = None
                for sep in [';', ',', '\t']:
                    try:
                        d = pd.read_csv(p, sep=sep, decimal=',', encoding='latin1', dtype=str)
                        if len(d.columns) > 5:
                            break
                    except Exception:
                        continue
                if d is None:
                    print(f'  Aviso: nao foi possivel ler {p.name} como CSV.')
                    continue
            elif ext in ('.xls', '.xlsx'):
                try:
                    d = pd.read_excel(p, dtype=str)
                except Exception:
                    d = pd.read_excel(p, dtype=str, engine='openpyxl')
            else:
                print(f'  Aviso: extensao nao suportada em {p.name}, ignorando.')
                continue

            d.columns = [c.strip().upper() for c in d.columns]
            d['_arquivo'] = p.name
            dfs.append(d)
            print(f'  Lido: {p.name} ({len(d)} linhas)')
        except Exception as e:
            print(f'  Aviso: erro ao ler {p.name}: {e}')

    if not dfs:
        raise ValueError('Nenhum arquivo valido encontrado.')

    df = pd.concat(dfs, ignore_index=True)

    num_cols = ['VLITEM', 'VLCONTABIL', 'VLBASEICMS', 'PERCICMS', 'VLICMS',
                'VLDESCONTO', 'QTCONT', 'VLTOTAL']
    for col in num_cols:
        if col in df.columns:
            df[col] = (
                pd.to_numeric(
                    df[col].astype(str).str.replace(',', '.').str.strip(),
                    errors='coerce'
                ).fillna(0)
            )

    for col in ['CODFISCAL', 'SITTRIBUT', 'ORIGMERCTRIB', 'NCM', 'CODFILIAL']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.strip()

    if 'CODFISCAL' in df.columns:
        df['CODFISCAL'] = df['CODFISCAL'].str[:4]
    if 'SITTRIBUT' in df.columns:
        df['SITTRIBUT'] = df['SITTRIBUT'].str.zfill(2)
    if 'PERCICMS' in df.columns:
        df['PERCICMS'] = pd.to_numeric(df['PERCICMS'], errors='coerce').fillna(0).round(2)
    if 'VLCONTABIL' not in df.columns:
        df['VLCONTABIL'] = df['VLITEM'] if 'VLITEM' in df.columns else 0.0

    if 'DATA' in df.columns:
        df['_DATA'] = pd.to_datetime(df['DATA'], dayfirst=True, errors='coerce')
    else:
        df['_DATA'] = pd.NaT

    df['TIPO_OP'] = 'Intraestadual'
    if 'CODFISCAL' in df.columns:
        df.loc[df['CODFISCAL'].str.startswith('6'), 'TIPO_OP'] = 'Interestadual'

    return df


_MESES_NOMES = {
    'janeiro':1,'fevereiro':2,'marco':3,'marco':3,'abril':4,'maio':5,'junho':6,
    'julho':7,'agosto':8,'setembro':9,'outubro':10,'novembro':11,'dezembro':12
}

def _periodo_do_nome(nome):
    import re
    nome_l = nome.lower()
    for palavra, mes_num in _MESES_NOMES.items():
        if palavra in nome_l:
            anos = re.findall(r'\b(20\d{2})\b', nome)
            ano = int(anos[0]) if anos else datetime.datetime.now().year
            return f'{MESES_PT[mes_num]} {ano}'
    return None

def extrair_periodo(df, caminhos=None):
    if caminhos:
        for c in caminhos:
            p = _periodo_do_nome(Path(c).name)
            if p:
                return p
    if 'DATA' not in df.columns:
        return datetime.datetime.now().strftime('%B %Y')
    datas = pd.to_datetime(df['DATA'], dayfirst=True, errors='coerce').dropna()
    if datas.empty:
        return datetime.datetime.now().strftime('%B %Y')
    from collections import Counter
    contagem = Counter(zip(datas.dt.year, datas.dt.month))
    ano, mes = contagem.most_common(1)[0][0]
    return f'{MESES_PT[mes]} {ano}'


# -- Divergencias --------------------------------------------------------------
def calcular_div1(df):
    """DIV-1: Aliquota 21% -- atipica em TO (CRITICA)"""
    d1 = df[df['PERCICMS'] == 21.0].copy()
    if d1.empty:
        return d1
    d1['ICMS Correto (R$)'] = (d1['VLBASEICMS'] * ALIQ_INT_DF / 100).round(2)
    d1['Dif. ICMS (R$)'] = (d1['VLICMS'] - d1['ICMS Correto (R$)']).round(2)
    d1['Critica'] = d1['NCM'].apply(
        lambda n: f'Aliq 21% em NCM {n} -- TO nao prevê 21%. Aliq geral TO e 20% (RICMS/DF (Decreto 18.955/97) + Lei 7.326/2023). Verificar supérfluo (25%) ou padrao (20%).'
    )
    return d1


def calcular_div2(intra):
    """DIV-2: Aliquota 12% intraestadual com CST 00 (ATENCAO)"""
    d2 = intra[(intra['PERCICMS'] == 12.0) & (intra['SITTRIBUT'] == '00')].copy()
    if d2.empty:
        return d2
    d2['ICMS Correto (R$)'] = (d2['VLBASEICMS'] * ALIQ_INT_DF / 100).round(2)
    d2['Dif. ICMS (R$)'] = (d2['VLICMS'] - d2['ICMS Correto (R$)']).round(2)
    d2['Critica'] = d2['NCM'].apply(
        lambda n: f'Aliq 12% intraestadual CST 00 em NCM {n} -- sem reducao formal no EFD (TO). Verificar convenio CONFAZ ou Lei TO.'
    )
    return d2


def calcular_div3(intra):
    """DIV-3: Base reduzida com CST 00 intraestadual (ATENCAO)"""
    cand = intra[
        (intra['SITTRIBUT'] == '00') &
        (intra['VLCONTABIL'] > 0) &
        (intra['VLBASEICMS'] > 0)
    ].copy()
    if cand.empty:
        return cand
    cand['_ratio'] = cand['VLBASEICMS'] / cand['VLCONTABIL']
    d3 = cand[cand['_ratio'] < 0.95].copy()
    if d3.empty:
        return d3
    d3['reducao_pct'] = ((1 - d3['_ratio']) * 100).round(2)
    d3['ICMS Correto (R$)'] = (d3['VLCONTABIL'] * d3['PERCICMS'] / 100).round(2)
    d3['Dif. ICMS (R$)'] = (d3['VLICMS'] - d3['ICMS Correto (R$)']).round(2)
    d3['Critica'] = d3['reducao_pct'].apply(
        lambda r: f'Base reduzida ~{r:.1f}% com CST 00 -- SPED exige CST 20 para base reduzida por convenio'
    )
    d3 = d3.drop(columns=['_ratio', 'reducao_pct'], errors='ignore')
    return d3


def calcular_div4(intra):
    """DIV-4: CST 20 -- verificar percentual de reducao (ATENCAO)"""
    d4 = intra[(intra['SITTRIBUT'] == '20') & (intra['VLCONTABIL'] > 0)].copy()
    if d4.empty:
        d4['ICMS Correto (R$)'] = pd.Series(dtype=float)
        d4['Dif. ICMS (R$)'] = pd.Series(dtype=float)
        d4['Critica'] = pd.Series(dtype=str)
        return d4
    d4['_ratio'] = (d4['VLBASEICMS'] / d4['VLCONTABIL'].replace(0, 1)).round(4)
    d4['reducao_pct'] = ((1 - d4['_ratio']) * 100).round(2)
    d4['ICMS Correto (R$)'] = 0.0
    d4['Dif. ICMS (R$)'] = 0.0
    d4['Critica'] = d4['reducao_pct'].apply(
        lambda r: f'CST 20 com reducao de {r:.1f}% -- confirmar convenio CONFAZ/DF e percentual correto.'
    )
    d4 = d4.drop(columns=['_ratio', 'reducao_pct'], errors='ignore')
    return d4


def calcular_div5(df):
    """DIV-5: CST 90 operacional (VERIFICAR)"""
    d5 = df[df['SITTRIBUT'] == '90'].copy()
    if d5.empty:
        d5['ICMS Correto (R$)'] = pd.Series(dtype=float)
        d5['Dif. ICMS (R$)'] = pd.Series(dtype=float)
        d5['Critica'] = pd.Series(dtype=str)
        return d5
    d5['ICMS Correto (R$)'] = 0.0
    d5['Dif. ICMS (R$)'] = 0.0
    d5['Critica'] = d5['CODFISCAL'].apply(
        lambda c: f'CST 90 em CFOP {c} -- verificar se ha tributacao pendente ou se CST e adequado'
    )
    return d5


def calcular_div6(inter):
    """DIV-6: CFOP 6xxx + Aliquota 4% + Origem nacional (CRITICA)"""
    if inter.empty:
        return inter
    d6 = inter[
        (inter['PERCICMS'] == 4.0) &
        (~inter['ORIGMERCTRIB'].isin(['1', '2', '6'])) &
        (inter['VLBASEICMS'] > 0)
    ].copy()
    if d6.empty:
        d6['ICMS Correto (R$)'] = pd.Series(dtype=float)
        d6['Dif. ICMS (R$)'] = pd.Series(dtype=float)
        d6['Critica'] = pd.Series(dtype=str)
        return d6
    d6['ICMS Correto (R$)'] = (d6['VLBASEICMS'] * 0.12).round(2)
    d6['Dif. ICMS (R$)'] = (d6['VLICMS'] - d6['ICMS Correto (R$)']).round(2)
    d6['Critica'] = d6['ORIGMERCTRIB'].apply(
        lambda o: f'Orig {o} (nacional) com aliq 4% interestadual -- Res. 13/2012 restringe 4% a importados. ICMS recolhido a menor.'
    )
    return d6


def calcular_div7(df):
    """DIV-7: CST 40/41 -- verificar convenio (VERIFICAR)"""
    d7 = df[df['SITTRIBUT'].isin(['40', '41'])].copy()
    if d7.empty:
        d7['ICMS Correto (R$)'] = pd.Series(dtype=float)
        d7['Dif. ICMS (R$)'] = pd.Series(dtype=float)
        d7['Critica'] = pd.Series(dtype=str)
        return d7
    d7['ICMS Correto (R$)'] = 0.0
    d7['Dif. ICMS (R$)'] = 0.0
    d7['Critica'] = d7.apply(
        lambda r: f'CST {r["SITTRIBUT"]} em NCM {r["NCM"]} -- confirmar convenio CONFAZ ou Lei TO que ampara a isencao',
        axis=1
    )
    return d7


def calcular_div8(intra):
    """DIV-8: Saidas intraestaduais com aliquota 18% apos 01/01/2024 -- TO especifico (ATENCAO)"""
    if intra.empty or '_DATA' not in intra.columns:
        return pd.DataFrame(columns=COLS_DIV)
    mask = (
        (intra['SITTRIBUT'] == '00') &
        (intra['PERCICMS'] == ALIQ_ANT_DF) &
        (intra['_DATA'].notna()) &
        (intra['_DATA'].dt.date >= DATA_MUDANCA)
    )
    d8 = intra[mask].copy()
    if d8.empty:
        return pd.DataFrame(columns=COLS_DIV)
    d8['ICMS Correto (R$)'] = (d8['VLBASEICMS'] * ALIQ_INT_DF / 100).round(2)
    d8['Dif. ICMS (R$)'] = (d8['VLICMS'] - d8['ICMS Correto (R$)']).round(2)
    d8['Critica'] = 'Saida intraestadual a 18% apos 01/01/2024 -- aliq. geral TO e 20% (RICMS/DF (Decreto 18.955/97) + Lei 7.326/2023). ICMS recolhido a menor.'
    return d8[[c for c in COLS_DIV if c in d8.columns]].copy()


# -- Base consolidada ----------------------------------------------------------
def calcular_base_consolidada(df):
    grp = df.groupby(['CODFISCAL', 'SITTRIBUT', 'PERCICMS']).agg(
        QTD_REGISTROS=('VLITEM', 'count'),
        QTD_NOTAS=('NUMNOTA', 'nunique'),
        VL_ITEM=('VLITEM', 'sum'),
        VL_BASE_ICMS=('VLBASEICMS', 'sum'),
        VL_ICMS=('VLICMS', 'sum'),
    ).reset_index()
    grp['BASE_ITEM_PCT'] = (
        grp['VL_BASE_ICMS'] / grp['VL_ITEM'].replace(0, 1) * 100
    ).round(2)
    grp['ALIQ_EFETIVA_PCT'] = (
        grp['VL_ICMS'] / grp['VL_BASE_ICMS'].replace(0, 1) * 100
    ).round(2)
    return grp


# -- Excel ---------------------------------------------------------------------
def prep_div(d):
    cols = [c for c in COLS_DIV if c in d.columns]
    return d[cols].copy()


def gerar_excel(df, divs, grp, periodo, caminho_excel):
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill, Font, Alignment
    from openpyxl.utils import get_column_letter

    divs_info = [
        ('DIV1', divs.get('DIV1'), 'Aliquota 21% atipica', 'CRITICA'),
        ('DIV2', divs.get('DIV2'), 'Aliquota 12% intraestadual CST 00', 'ATENCAO'),
        ('DIV3', divs.get('DIV3'), 'Base reduzida com CST 00', 'ATENCAO'),
        ('DIV4', divs.get('DIV4'), 'CST 20 -- verificar percentual', 'ATENCAO'),
        ('DIV5', divs.get('DIV5'), 'CST 90 operacional', 'VERIFICAR'),
        ('DIV6', divs.get('DIV6'), 'Interestadual 4% origem nacional', 'CRITICA'),
        ('DIV7', divs.get('DIV7'), 'CST 40/41 -- verificar convenio', 'VERIFICAR'),
        ('DIV8', divs.get('DIV8'), 'Saida 18% apos 01/01/2024 (TO)', 'ATENCAO'),
    ]

    resumo_rows = []
    for nome_aba, dados, desc, crit in divs_info:
        n = len(dados) if dados is not None and not dados.empty else 0
        vl = round(dados['VLCONTABIL'].sum(), 2) if dados is not None and not dados.empty and 'VLCONTABIL' in dados.columns else 0
        dif = round(dados['Dif. ICMS (R$)'].sum(), 2) if dados is not None and not dados.empty and 'Dif. ICMS (R$)' in dados.columns else 0
        resumo_rows.append({
            'DIV': nome_aba.replace('DIV', 'DIV-'),
            'Descricao': desc,
            'Criticidade': crit,
            'Qtd Registros': n,
            'Valor Envolvido (R$)': vl,
            'Dif. ICMS (R$)': dif
        })

    resumo_df = pd.DataFrame(resumo_rows)

    with pd.ExcelWriter(caminho_excel, engine='openpyxl') as writer:
        resumo_df.to_excel(writer, sheet_name='RESUMO', index=False, startrow=3)
        grp.to_excel(writer, sheet_name='BASE CONSOLIDADA', index=False, startrow=3)
        for nome_aba, dados, _, _ in divs_info:
            if dados is not None and not dados.empty:
                df_aba = prep_div(dados)
            else:
                df_aba = pd.DataFrame(columns=COLS_DIV[:5])
            df_aba.to_excel(writer, sheet_name=nome_aba, index=False, startrow=3)

    TAB_COLORS = {
        'RESUMO': '1F4E79',
        'BASE CONSOLIDADA': '2E75B6',
        'DIV1': 'C00000', 'DIV2': 'ED7D31', 'DIV3': 'ED7D31',
        'DIV4': 'ED7D31', 'DIV5': '375623', 'DIV6': 'C00000',
        'DIV7': '375623', 'DIV8': 'ED7D31',
    }
    DARK_FILL = PatternFill('solid', fgColor='1F4E79')
    WHITE_FONT = Font(bold=True, color='FFFFFF', size=11)
    LIGHT_FILL = PatternFill('solid', fgColor='F2F2F2')

    wb = load_workbook(caminho_excel)

    for sname in wb.sheetnames:
        ws = wb[sname]
        ws.sheet_properties.tabColor = TAB_COLORS.get(sname, '1F4E79')

        max_col = max(ws.max_column, 1)
        ws.merge_cells(f'A1:{get_column_letter(max_col)}1')
        t = ws['A1']
        if sname == 'RESUMO':
            t.value = f'ANALISE DE SAIDAS ICMS/DF -- {periodo.upper()}'
        else:
            t.value = f'{sname} -- Divergencias de Saidas ICMS/DF -- {periodo}'
        t.font = WHITE_FONT
        t.fill = DARK_FILL
        t.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 28

        for cell in ws[4]:
            cell.font = Font(bold=True, color='FFFFFF', size=10)
            cell.fill = DARK_FILL
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws.row_dimensions[4].height = 28

        if ws.max_row > 4:
            ws.auto_filter.ref = ws.dimensions

        for row in ws.iter_rows(min_row=5):
            fill = LIGHT_FILL if row[0].row % 2 == 0 else PatternFill('solid', fgColor='FFFFFF')
            for cell in row:
                try:
                    rgb = cell.fill.fgColor.rgb
                    if rgb not in ('FCE4D6', 'E2EFDA'):
                        cell.fill = fill
                except Exception:
                    cell.fill = fill

        dif_col = None
        for cell in ws[4]:
            if cell.value and 'Dif' in str(cell.value):
                dif_col = cell.column
                break
        if dif_col:
            for row in ws.iter_rows(min_row=5, min_col=dif_col, max_col=dif_col):
                for cell in row:
                    try:
                        v = float(cell.value or 0)
                        if v > 0:
                            cell.fill = PatternFill('solid', fgColor='E2EFDA')
                            cell.font = Font(bold=True, color='375623')
                        elif v < 0:
                            cell.fill = PatternFill('solid', fgColor='FCE4D6')
                            cell.font = Font(bold=True, color='C00000')
                    except Exception:
                        pass

        for col in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                try:
                    cell_len = len(str(cell.value or ''))
                    if cell_len > max_len:
                        max_len = cell_len
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = min(max(max_len + 2, 10), 50)

    wb.save(caminho_excel)


# -- Word ----------------------------------------------------------------------
def gerar_word(df, divs_dict, periodo, caminho_word):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    h = doc.add_heading(f'Analise de Saidas ICMS/DF -- {periodo}', 0)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    intra = df[df['TIPO_OP'] == 'Intraestadual']
    inter = df[df['TIPO_OP'] == 'Interestadual']

    h1 = doc.add_heading('1. Visao Geral', 1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    tbl = doc.add_table(rows=8, cols=2)
    tbl.style = 'Table Grid'
    aliq_ef = (df['VLICMS'].sum() / df['VLBASEICMS'].sum() * 100) if df['VLBASEICMS'].sum() > 0 else 0
    dados_vg = [
        ('Total de Registros', f'{len(df):,}'),
        ('Valor Total Saidas (R$)', f'{df["VLCONTABIL"].sum():,.2f}'),
        ('Base ICMS Total (R$)', f'{df["VLBASEICMS"].sum():,.2f}'),
        ('ICMS Total (R$)', f'{df["VLICMS"].sum():,.2f}'),
        ('Intraestadual (registros)', f'{len(intra):,}'),
        ('Interestadual (registros)', f'{len(inter):,}'),
        ('Aliquota Efetiva Media', f'{aliq_ef:.2f}%'),
        ('Periodo', periodo),
    ]
    for i, (label, value) in enumerate(dados_vg):
        row = tbl.rows[i].cells
        row[0].text = label; row[1].text = str(value)

    doc.add_paragraph()

    infos_div = [
        ('DIV-1', 'Aliquota 21% -- Atipica em TO', 'CRITICA',
         'A aliquota de 21% nao esta prevista no RICMS/DF vigente. A aliquota geral do Distrito Federal e de 20% desde 01/01/2024 (RICMS/DF (Decreto 18.955/97) + Lei 7.326/2023). Risco bidirecional: se deveria ser 20%, ha recolhimento a maior (recuperavel); se supérfluo (25%), ha recolhimento a menor.',
         'Verificar a tabela de aliquotas do RICMS/DF para cada NCM; se padrao, corrigir para 20%; se supérfluo, verificar se 25% se aplica no TO.'),
        ('DIV-2', 'Aliquota 12% Intraestadual com CST 00', 'ATENCAO',
         'Aliquota de 12% em operacao interna sem CST 20 significa que a reducao de base nao esta formalmente declarada no EFD. Verificar se ha convenio CONFAZ ou Lei TO especifica para o NCM. A SEFAZ-DF (www.fazenda.df.gov.br) disponibiliza a legislacao vigente.',
         'Identificar o convenio de amparo para cada NCM; parametrizar o ERP para emitir CST 20; retificar o EFD dos periodos com a declaracao correta.'),
        ('DIV-3', 'Base de Calculo Reduzida com CST 00', 'ATENCAO',
         'Reducao de base de calculo sem declaracao de CST 20 no EFD e incorreto. O SPED Fiscal cruza o CST com a relacao base/valor e pode gerar pendencia automatica.',
         'Identificar o convenio; parametrizar o ERP para CST 20; retificar o EFD. Verificar se ha estorno proporcional de credito quando exigido pelo convenio.'),
        ('DIV-4', 'CST 20 -- Verificar Percentual de Reducao', 'ATENCAO',
         'CST 20 esta formalmente correto, mas o percentual de reducao aplicado deve corresponder exatamente ao convenio de amparo. Percentuais incorretos geram ICMS a maior ou a menor.',
         'Confirmar o percentual de reducao no convenio especifico para cada NCM; verificar se ha exigencia de estorno proporcional de credito; ajustar a parametrizacao do ERP.'),
        ('DIV-5', 'CST 90 em CFOP Operacional', 'VERIFICAR',
         'CST 90 nao deve ser usado como codigo generico. CFOPs como 5927 e 5949 normalmente tem ICMS zero, mas itens com natureza diferente podem ter tributacao pendente.',
         'Revisar cada item com CST 90; identificar se ha tributacao pendente; ajustar CST conforme a natureza real da operacao.'),
        ('DIV-6', 'Interestadual 4% em Origem Nacional', 'CRITICA',
         'A Resolucao do Senado 13/2012 restringe a aliquota de 4% exclusivamente a mercadorias importadas (origem 1, 2 ou 6). Mercadorias de origem nacional tributadas a 4% em operacoes interestaduais representam ICMS recolhido a menor.',
         'Emitir complemento de ICMS para os periodos identificados; corrigir a parametrizacao do ERP; acompanhar possiveis notificacoes do estado destinatario.'),
        ('DIV-7', 'CST 40/41 -- Verificar Convenio', 'VERIFICAR',
         'Isencoes e nao tributacoes devem ter amparo em convenio CONFAZ ou lei estadual especifica. Produtos como livros tem imunidade constitucional; demais precisam de verificacao individual. Consultar www.fazenda.df.gov.br para a lista de beneficios vigentes no TO.',
         'Listar os convenios que amparam cada NCM isento; arquivar a documentacao no dossie fiscal; verificar se ha NCMs sem amparo identificado.'),
        ('DIV-8', 'Saida Intraestadual a 18% apos 01/01/2024 (Especifico TO)', 'ATENCAO',
         'A aliquota geral interna do Distrito Federal foi alterada de 18% para 20% em 01/01/2024 pela RICMS/DF (Decreto 18.955/97) + Lei 7.326/2023. Saidas intraestaduais emitidas apos essa data ainda a 18% indicam ERP com aliquota desatualizada, gerando ICMS recolhido a menor e risco de autuacao na SEFAZ-DF.',
         'Atualizar imediatamente a aliquota interna do TO no ERP para 20%; verificar o estoque de NFs emitidas a menor e avaliar emissao de nota complementar; consultar o departamento juridico/tributario sobre a prescricao e risco de autuacao.'),
    ]

    cor_critica = {
        'CRITICA': RGBColor(0xC0, 0x00, 0x00),
        'ATENCAO': RGBColor(0xC5, 0x5A, 0x11),
        'VERIFICAR': RGBColor(0x37, 0x56, 0x23)
    }

    h2_header = doc.add_heading('2. Divergencias Identificadas', 1)
    h2_header.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    for div_id, titulo, crit, narrativa, acao in infos_div:
        chave = div_id.replace('-', '')
        dados = divs_dict.get(chave)
        n_regs = len(dados) if dados is not None and not dados.empty else 0

        h2 = doc.add_heading(f'{div_id} -- {titulo}', 2)
        h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        p = doc.add_paragraph()
        run = p.add_run(f'  {crit} -- {n_regs} registros identificados')
        run.bold = True
        run.font.color.rgb = cor_critica.get(crit, RGBColor(0x00, 0x00, 0x00))

        doc.add_paragraph(narrativa)

        if dados is not None and not dados.empty:
            top_n = min(5, len(dados))
            if 'VLCONTABIL' in dados.columns:
                top = dados.nlargest(top_n, 'VLCONTABIL')
            else:
                top = dados.head(top_n)
            doc.add_paragraph('Principais casos:')
            for _, row in top.iterrows():
                nf = row.get('NUMNOTA', '--')
                vl = float(row.get('VLCONTABIL', row.get('VLITEM', 0)))
                desc = str(row.get('DESCRICAO', ''))[:50]
                doc.add_paragraph(f'NF {nf} -- R$ {vl:,.2f} -- {desc}', style='List Bullet')

        p_acao = doc.add_paragraph()
        p_acao.add_run('Acao recomendada: ').bold = True
        p_acao.add_run(acao)
        doc.add_paragraph()

    h_res = doc.add_heading('3. Resumo e Prioridades', 1)
    h_res.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = 'Table Grid'
    hdr = tbl2.rows[0].cells
    for i, h in enumerate(['DIV', 'Criticidade', 'Registros', 'Dif. ICMS (R$)']):
        hdr[i].text = h

    for div_id, titulo, crit, _, _ in infos_div:
        chave = div_id.replace('-', '')
        dados = divs_dict.get(chave)
        n = len(dados) if dados is not None and not dados.empty else 0
        dif = 0.0
        if dados is not None and not dados.empty and 'Dif. ICMS (R$)' in dados.columns:
            dif = round(dados['Dif. ICMS (R$)'].sum(), 2)
        row_cells = tbl2.add_row().cells
        row_cells[0].text = div_id; row_cells[1].text = crit
        row_cells[2].text = str(n); row_cells[3].text = f'R$ {dif:,.2f}'

    doc.add_paragraph()
    p = doc.add_paragraph(
        f'Gerado automaticamente em {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")} '
        f'-- Especialista em Apuracao de ICMS'
    )
    if p.runs:
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(str(caminho_word))


# -- main ----------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print('Uso: python analisar_saidas_df.py <arquivo1.csv> [arquivo2.xls] ...')
        sys.exit(1)

    caminhos = sys.argv[1:]
    print(f'Carregando {len(caminhos)} arquivo(s)...')
    df = carregar_dados(caminhos)
    print(f'  {len(df)} registros | {df["CODFILIAL"].nunique() if "CODFILIAL" in df.columns else "?"} filiais')

    periodo = extrair_periodo(df, caminhos)
    print(f'  Periodo: {periodo}')

    intra = df[df['TIPO_OP'] == 'Intraestadual']
    inter = df[df['TIPO_OP'] == 'Interestadual']

    print('Calculando divergencias...')
    d1 = calcular_div1(df)
    d2 = calcular_div2(intra)
    d3 = calcular_div3(intra)
    d4 = calcular_div4(intra)
    d5 = calcular_div5(df)
    d6 = calcular_div6(inter)
    d7 = calcular_div7(df)
    d8 = calcular_div8(intra)

    divs = {
        'DIV1': d1, 'DIV2': d2, 'DIV3': d3, 'DIV4': d4,
        'DIV5': d5, 'DIV6': d6, 'DIV7': d7, 'DIV8': d8
    }
    for k, v in divs.items():
        print(f'  {k}: {len(v) if v is not None else 0} registros')

    grp = calcular_base_consolidada(df)

    caminho_excel = PASTA_SAIDA / f'Analise Saidas ICMS DF - {periodo}.xlsx'
    caminho_word  = PASTA_SAIDA / f'Analise Saidas ICMS DF - {periodo}.docx'

    print(f'Gerando Excel: {caminho_excel.name}')
    gerar_excel(df, divs, grp, periodo, caminho_excel)

    print(f'Gerando Word: {caminho_word.name}')
    gerar_word(df, divs, periodo, caminho_word)

    print('Concluido!')


if __name__ == '__main__':
    main()
