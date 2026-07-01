"""
analisar_entradas_es.py
Analise de Entradas ICMS/ES (Espirito Santo) -- gera Excel (.xlsx) e Word (.docx)
Uso: python analisar_entradas_es.py <arquivo1.xls> [arquivo2.csv] ...
"""

from __future__ import annotations
import sys
import datetime
from pathlib import Path

import pandas as pd
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor

# -- Constantes ----------------------------------------------------------------
PASTA_SAIDA = Path(__file__).parent
ALIQ_INT_ES = 17.0               # RICMS/ES (Decreto 1.090-R/2002) -- aliq. geral ES

MESES_PT = {
    1: 'Janeiro', 2: 'Fevereiro', 3: 'Marco', 4: 'Abril',
    5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
    9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
}

COLS_DIV = [
    'CODFILIAL', 'NUMNOTA', 'DATA', 'CST', 'CODFISCAL', 'ORIGMERCTRIB',
    'CODPROD', 'DESCRICAO', 'NCM', 'VLITEM', 'VLBASEICMS',
    'PERCICMS', 'VLICMS', 'ICMS Correto (R$)', 'Dif. ICMS (R$)', 'Critica', '_arquivo'
]

CFOP_DESC = {
    '2923': 'complemento de preco ou bonificacao financeira',
    '2557': 'frete CIF -- integra base ICMS da mercadoria',
    '2910': 'bonificacao/brinde -- CST correto: 40 ou 41',
    '2551': 'ativo imobilizado -- verificar se e frete ou instalacao',
    '2556': 'retorno de servico',
    '2949': 'materiais graficos/amostras',
}


# -- Leitura -------------------------------------------------------------------
def carregar_dados(caminhos):
    dfs = []
    for c in caminhos:
        p = Path(c)
        try:
            if p.suffix.lower() == '.csv':
                d = None
                for sep in [';', ',', '\t']:
                    try:
                        tmp = pd.read_csv(p, sep=sep, decimal=',', encoding='latin1', dtype=str)
                        if len(tmp.columns) > 5:
                            d = tmp
                            break
                    except Exception:
                        continue
                if d is None:
                    print(f'  Aviso: nao foi possivel ler {p.name} como CSV')
                    continue
            else:
                try:
                    d = pd.read_excel(p, dtype=str)
                except Exception:
                    d = pd.read_excel(p, dtype=str, engine='openpyxl')
            d.columns = [x.strip().upper() for x in d.columns]
            d['_arquivo'] = p.name
            dfs.append(d)
            print(f'  Lido: {p.name} ({len(d):,} linhas)')
        except Exception as e:
            print(f'  Aviso: {p.name}: {e}')

    if not dfs:
        raise ValueError('Nenhum arquivo pode ser lido. Verifique o formato (XLS/XLSX/CSV).')

    df = pd.concat(dfs, ignore_index=True)

    num_cols = ['VLITEM', 'VLCONTABIL', 'VLBASEICMS', 'PERCICMS', 'VLICMS',
                'VLDESCONTO', 'QTCONT', 'VLPIS', 'VLCOFINS', 'VLBASEPISCOFINS',
                'PERPIS', 'PERCOFINS', 'VLNAOTRIBUTADO']
    for col in num_cols:
        if col in df.columns:
            df[col] = (
                pd.to_numeric(
                    df[col].astype(str).str.replace(',', '.').str.strip(),
                    errors='coerce'
                ).fillna(0)
            )

    for col in ['CODFISCAL', 'CST', 'ORIGMERCTRIB', 'NCM', 'CODFILIAL']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.strip()

    df['CODFISCAL'] = df['CODFISCAL'].str[:4]
    if 'CST' in df.columns:
        df['CST'] = df['CST'].str.zfill(2)
    df['PERCICMS'] = pd.to_numeric(df['PERCICMS'], errors='coerce').fillna(0).round(2)

    if 'VLCONTABIL' not in df.columns:
        df['VLCONTABIL'] = df['VLITEM'] if 'VLITEM' in df.columns else 0.0

    if 'DATA' in df.columns:
        df['_DATA'] = pd.to_datetime(df['DATA'], dayfirst=True, errors='coerce')
    else:
        df['_DATA'] = pd.NaT

    df['TIPO_OP'] = 'Intraestadual'
    mask_inter = (
        df['CODFISCAL'].str.startswith('2') | df['CODFISCAL'].str.startswith('3')
    )
    df.loc[mask_inter, 'TIPO_OP'] = 'Interestadual'

    return df


_MESES_NOMES = {
    'janeiro':1,'fevereiro':2,'marco':3,'marco':3,'abril':4,'maio':5,'junho':6,
    'julho':7,'agosto':8,'setembro':9,'outubro':10,'novembro':11,'dezembro':12
}

def _periodo_do_nome(nome):
    import re
    for palavra, mes_num in _MESES_NOMES.items():
        if palavra in nome.lower():
            anos = re.findall(r'\b(20\d{2})\b', nome)
            ano = int(anos[0]) if anos else datetime.datetime.now().year
            return f'{MESES_PT[mes_num]} {ano}'
    return None

def extrair_periodo(df, caminhos=None):
    if caminhos:
        for c in caminhos:
            p = _periodo_do_nome(Path(c).name)
            if p:
                return p
    if 'DATA' not in df.columns:
        return datetime.datetime.now().strftime('%B %Y')
    datas = pd.to_datetime(df['DATA'], dayfirst=True, errors='coerce').dropna()
    if datas.empty:
        return datetime.datetime.now().strftime('%B %Y')
    from collections import Counter
    contagem = Counter(zip(datas.dt.year, datas.dt.month))
    ano, mes = contagem.most_common(1)[0][0]
    return f'{MESES_PT[mes]} {ano}'


# -- Divergencias --------------------------------------------------------------
def _selecionar_cols(df):
    return [c for c in COLS_DIV if c in df.columns]


def calcular_div1(df, inter):
    """DIV-1: Importada (Orig 1/2/6) + CFOP 2xxx + aliquota != 4% (CRITICA)"""
    d1 = inter[
        inter['ORIGMERCTRIB'].isin(['1', '2', '6']) &
        (inter['CST'] == '00') &
        (inter['PERCICMS'] != 4.0) &
        (inter['PERCICMS'] > 0)
    ].copy()
    d1['ICMS Correto (R$)'] = (d1['VLBASEICMS'] * 0.04).round(2)
    d1['Dif. ICMS (R$)'] = (d1['VLICMS'] - d1['ICMS Correto (R$)']).round(2)
    d1['Critica'] = d1.apply(
        lambda r: f'Orig {r["ORIGMERCTRIB"]} a {r["PERCICMS"]}% -- deveria ser 4% (Res. Senado 13/2012)',
        axis=1
    )
    return d1[_selecionar_cols(d1)]


def calcular_div2(inter):
    """DIV-2: CST 41 interestadual (ATENCAO)"""
    d2 = inter[inter['CST'] == '41'].copy()
    d2['ICMS Correto (R$)'] = 0.0
    d2['Dif. ICMS (R$)'] = 0.0
    d2['Critica'] = 'CST 41 interestadual -- solicitar documentacao do convenio/isencao aplicada'
    return d2[_selecionar_cols(d2)]


def calcular_div3(inter):
    """DIV-3: CST 90 interestadual (ATENCAO)"""
    d3 = inter[inter['CST'] == '90'].copy()
    d3['ICMS Correto (R$)'] = 0.0
    d3['Dif. ICMS (R$)'] = 0.0
    d3['Critica'] = d3['CODFISCAL'].apply(
        lambda c: f'CFOP {c} + CST 90: {CFOP_DESC.get(c, "verificar natureza da operacao")}'
    )
    return d3[_selecionar_cols(d3)]


def calcular_div4(inter):
    """DIV-4: Base reduzida com CST 00 interestadual (ATENCAO)"""
    cand = inter[
        (inter['CST'] == '00') &
        (inter['VLCONTABIL'] > 0) &
        (inter['VLBASEICMS'] > 0)
    ].copy()
    if cand.empty:
        return pd.DataFrame(columns=COLS_DIV)
    cand['_ratio'] = cand['VLBASEICMS'] / cand['VLCONTABIL']
    d4 = cand[cand['_ratio'] < 0.95].copy()
    if d4.empty:
        return pd.DataFrame(columns=COLS_DIV)
    d4['reducao_pct'] = ((1 - d4['_ratio']) * 100).round(2)
    d4['ICMS Correto (R$)'] = 0.0
    d4['Dif. ICMS (R$)'] = 0.0
    d4['Critica'] = d4['reducao_pct'].apply(
        lambda r: f'Base reduzida ~{r:.1f}% com CST 00 -- EFD exige CST 20 para base reduzida por convenio'
    )
    return d4[_selecionar_cols(d4)]


def calcular_div5(inter):
    """DIV-5: Aliquota atipica interestadual (VERIFICAR)"""
    VALIDAS = {0.0, 4.0, 7.0, 12.0}
    d5 = inter[
        (inter['CST'] == '00') &
        (~inter['PERCICMS'].isin(VALIDAS)) &
        (inter['PERCICMS'] > 0) &
        (~inter['ORIGMERCTRIB'].isin(['1', '2', '6']))
    ].copy()
    d5['ICMS Correto (R$)'] = 0.0
    d5['Dif. ICMS (R$)'] = 0.0
    d5['Critica'] = d5['PERCICMS'].apply(
        lambda a: f'Aliquota {a}% nao prevista na tabela interestadual (esperado: 7% ou 12%)'
    )
    return d5[_selecionar_cols(d5)]




# -- Excel ---------------------------------------------------------------------
def _thin_border():
    thin = Side(style='thin')
    return Border(left=thin, right=thin, top=thin, bottom=thin)

def _escrever_titulo(ws, texto, cor_hex, n_cols):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=n_cols)
    c = ws.cell(row=1, column=1, value=texto)
    c.font = Font(bold=True, size=14, color='FFFFFF')
    c.fill = PatternFill('solid', fgColor=cor_hex)
    c.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 22

def _escrever_cabecalho(ws, colunas, cor_hex, linha=4):
    for j, col in enumerate(colunas, 1):
        c = ws.cell(row=linha, column=j, value=col)
        c.font = Font(bold=True, color='FFFFFF', size=10)
        c.fill = PatternFill('solid', fgColor=cor_hex)
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        c.border = _thin_border()
    ws.row_dimensions[linha].height = 30

def _colorir_dif(ws, col_idx, n_linhas, linha_dados=5):
    fill_verde = PatternFill('solid', fgColor='E2EFDA')
    fill_verm  = PatternFill('solid', fgColor='FCE4D6')
    font_verde = Font(bold=True, color='375623')
    font_verm  = Font(bold=True, color='C00000')
    for row in range(linha_dados, linha_dados + n_linhas):
        c = ws.cell(row=row, column=col_idx)
        try:
            v = float(c.value) if c.value is not None else 0
        except (TypeError, ValueError):
            v = 0
        if v > 0:
            c.fill = fill_verde; c.font = font_verde
        elif v < 0:
            c.fill = fill_verm; c.font = font_verm

def _escrever_aba_div(wb, nome_aba, df_div, cor_hex, titulo):
    ws = wb.create_sheet(nome_aba)
    if df_div is None or df_div.empty:
        _escrever_titulo(ws, titulo, cor_hex, 5)
        ws.cell(row=3, column=1, value='Nenhuma divergencia encontrada nesta categoria.')
        ws.cell(row=3, column=1).font = Font(italic=True, color='555555')
        return
    colunas = [c for c in COLS_DIV if c in df_div.columns]
    n_cols = len(colunas)
    _escrever_titulo(ws, titulo, cor_hex, max(n_cols, 5))
    _escrever_cabecalho(ws, colunas, cor_hex, linha=4)
    df_out = df_div[colunas].copy()
    for i, (_, row) in enumerate(df_out.iterrows(), start=5):
        for j, col in enumerate(colunas, 1):
            val = row[col]
            if pd.isna(val): val = ''
            c = ws.cell(row=i, column=j, value=val)
            c.border = _thin_border()
            c.alignment = Alignment(vertical='center', wrap_text=(col == 'Critica'))
            if col in ('VLITEM', 'VLBASEICMS', 'VLICMS', 'ICMS Correto (R$)', 'Dif. ICMS (R$)', 'PERCICMS'):
                c.number_format = '#,##0.00'
    if 'Dif. ICMS (R$)' in colunas:
        dif_col_idx = colunas.index('Dif. ICMS (R$)') + 1
        _colorir_dif(ws, dif_col_idx, len(df_out), linha_dados=5)
    row_tot = 5 + len(df_out)
    ws.cell(row=row_tot, column=1, value='TOTAL').font = Font(bold=True)
    for j, col in enumerate(colunas, 1):
        if col in ('VLITEM', 'VLBASEICMS', 'VLICMS', 'ICMS Correto (R$)', 'Dif. ICMS (R$)'):
            try:
                total = df_out[col].sum()
                c = ws.cell(row=row_tot, column=j, value=round(total, 2))
                c.font = Font(bold=True); c.number_format = '#,##0.00'; c.border = _thin_border()
            except Exception:
                pass
    ws.auto_filter.ref = f'A4:{get_column_letter(n_cols)}{row_tot}'
    for j, col in enumerate(colunas, 1):
        ltr = get_column_letter(j)
        if col in ('DESCRICAO', 'Critica', 'PARTICIPANTE'): ws.column_dimensions[ltr].width = 45
        elif col in ('VLITEM', 'VLBASEICMS', 'VLICMS', 'ICMS Correto (R$)', 'Dif. ICMS (R$)'): ws.column_dimensions[ltr].width = 18
        elif col in ('DATA', 'NUMNOTA', 'CODFISCAL', 'CST', 'ORIGMERCTRIB', 'NCM', 'PERCICMS'): ws.column_dimensions[ltr].width = 13
        else: ws.column_dimensions[ltr].width = 12

def _aba_resumo(wb, df, divs, periodo):
    ws = wb.create_sheet('RESUMO', 0)
    COR = '1F4E79'
    _escrever_titulo(ws, f'Analise de Entradas ICMS/ES -- {periodo}', COR, 4)
    inter = df[df['TIPO_OP'] == 'Interestadual']
    intra = df[df['TIPO_OP'] == 'Intraestadual']
    dados = [
        ('Total de Registros', f'{len(df):,}', '', ''),
        ('Intraestadual (registros)', f'{len(intra):,}', 'Interestadual (registros)', f'{len(inter):,}'),
        ('Valor Total Entradas (R$)', f'{df["VLCONTABIL"].sum():,.2f}', 'Base ICMS Total (R$)', f'{df["VLBASEICMS"].sum():,.2f}'),
        ('ICMS Total (R$)', f'{df["VLICMS"].sum():,.2f}', 'Periodo', periodo),
        ('', '', '', ''),
    ]
    r = 3
    ws.merge_cells(f'A{r}:D{r}')
    c = ws.cell(row=r, column=1, value='VISAO GERAL')
    c.font = Font(bold=True, color='FFFFFF', size=11)
    c.fill = PatternFill('solid', fgColor=COR)
    c.alignment = Alignment(horizontal='center')
    for row_data in dados:
        r += 1
        for j, val in enumerate(row_data, 1):
            cell = ws.cell(row=r, column=j, value=val)
            if j % 2 == 1 and val:
                cell.font = Font(bold=True); cell.fill = PatternFill('solid', fgColor='D6E4F0')
    r += 2
    ws.merge_cells(f'A{r}:D{r}')
    c = ws.cell(row=r, column=1, value='RESUMO DE DIVERGENCIAS')
    c.font = Font(bold=True, color='FFFFFF', size=11)
    c.fill = PatternFill('solid', fgColor=COR)
    c.alignment = Alignment(horizontal='center')
    r += 1
    cabecalhos = ['DIV', 'Criticidade', 'Registros', 'Dif. ICMS (R$)']
    cores_crit = {'CRITICA': 'C00000', 'ATENCAO': 'ED7D31', 'VERIFICAR': '375623'}
    for j, h in enumerate(cabecalhos, 1):
        c = ws.cell(row=r, column=j, value=h)
        c.font = Font(bold=True, color='FFFFFF'); c.fill = PatternFill('solid', fgColor=COR)
        c.alignment = Alignment(horizontal='center'); c.border = _thin_border()
    infos_resumo = [
        ('DIV1', 'CRITICA'), ('DIV2', 'ATENCAO'),
        ('DIV3', 'ATENCAO'), ('DIV4', 'ATENCAO'), ('DIV5', 'VERIFICAR'),
    ]
    for div_id, crit in infos_resumo:
        r += 1
        dados_div = divs.get(div_id)
        n = len(dados_div) if dados_div is not None and not dados_div.empty else 0
        dif = round(dados_div['Dif. ICMS (R$)'].sum(), 2) if dados_div is not None and not dados_div.empty and 'Dif. ICMS (R$)' in dados_div.columns else 0.0
        vals = [div_id, crit, n, dif]
        for j, val in enumerate(vals, 1):
            c = ws.cell(row=r, column=j, value=val); c.border = _thin_border()
            c.alignment = Alignment(horizontal='center' if j != 2 else 'left')
            if j == 2: c.font = Font(bold=True, color=cores_crit.get(crit, '000000'))
            if j == 4: c.number_format = '#,##0.00'
    for ltr, w in [('A', 12), ('B', 14), ('C', 14), ('D', 18)]:
        ws.column_dimensions[ltr].width = w

def _aba_sintese(wb, divs, periodo):
    ws = wb.create_sheet('SINTESE RISCOS', 1)
    COR = '1F4E79'
    _escrever_titulo(ws, f'Sintese de Riscos -- Entradas ICMS/ES -- {periodo}', COR, 7)
    colunas = ['DIV', 'GRUPO', 'VALOR ENVOLVIDO (R$)', 'REGISTROS', 'IMP. ICMS (R$)', 'NATUREZA PROVAVEL', 'RISCO']
    _escrever_cabecalho(ws, colunas, COR, linha=4)
    grupos_info = {
        'DIV1': {'grupo': 'Importada Aliq. > 4%', 'natureza': 'Credito indevido -- Res. Senado 13/2012', 'risco': 'CRITICA', 'cor': 'C00000'},
        'DIV2': {'grupo': 'CST 41 Interestadual', 'natureza': 'Isencao sem amparo ou convenio nao documentado', 'risco': 'ATENCAO', 'cor': 'ED7D31'},
        'DIV3': {'grupo': 'CST 90 Interestadual', 'natureza': 'Natureza da operacao nao identificada', 'risco': 'ATENCAO', 'cor': 'ED7D31'},
        'DIV4': {'grupo': 'Base Reduzida sem CST 20', 'natureza': 'CST incorreto -- EFD exige CST 20', 'risco': 'ATENCAO', 'cor': 'ED7D31'},
        'DIV5': {'grupo': 'Aliquota Atipica Inter.', 'natureza': 'Erro de cadastro ou convenio nao documentado', 'risco': 'VERIFICAR', 'cor': '375623'},
    }
    r = 5
    for div_id, info in grupos_info.items():
        dados_div = divs.get(div_id)
        n = len(dados_div) if dados_div is not None and not dados_div.empty else 0
        vl = round(dados_div['VLCONTABIL'].sum(), 2) if dados_div is not None and not dados_div.empty and 'VLCONTABIL' in dados_div.columns else 0.0
        dif = round(dados_div['Dif. ICMS (R$)'].sum(), 2) if dados_div is not None and not dados_div.empty and 'Dif. ICMS (R$)' in dados_div.columns else 0.0
        vals = [div_id, info['grupo'], vl, n, dif, info['natureza'], info['risco']]
        for j, val in enumerate(vals, 1):
            c = ws.cell(row=r, column=j, value=val); c.border = _thin_border()
            c.alignment = Alignment(vertical='center', wrap_text=(j == 6), horizontal='center' if j in (1, 4, 7) else 'left')
            if j in (3, 5): c.number_format = '#,##0.00'
            if j == 7: c.font = Font(bold=True, color=info['cor'])
        r += 1
    for ltr, w in [('A', 10), ('B', 32), ('C', 22), ('D', 12), ('E', 18), ('F', 45), ('G', 12)]:
        ws.column_dimensions[ltr].width = w
    ws.auto_filter.ref = f'A4:G{r - 1}'

def gerar_excel(df, divs, periodo, caminho_excel):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    _aba_resumo(wb, df, divs, periodo)
    _aba_sintese(wb, divs, periodo)
    abas_div = [
        ('DIV1', 'C00000', 'DIV-1 -- Importada com Aliquota > 4% (CRITICA)'),
        ('DIV2', 'ED7D31', 'DIV-2 -- CST 41 Interestadual (ATENCAO)'),
        ('DIV3', 'ED7D31', 'DIV-3 -- CST 90 Interestadual (ATENCAO)'),
        ('DIV4', 'ED7D31', 'DIV-4 -- Base Reduzida sem CST 20 (ATENCAO)'),
        ('DIV5', '375623', 'DIV-5 -- Aliquota Atipica Interestadual (VERIFICAR)'),
    ]
    for nome, cor, titulo in abas_div:
        _escrever_aba_div(wb, nome, divs.get(nome), cor, titulo)
    wb.save(str(caminho_excel))


# -- Word ----------------------------------------------------------------------
def gerar_word(df, divs_dict, periodo, caminho_word):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(36); section.bottom_margin = Pt(36)
        section.left_margin = Pt(54); section.right_margin = Pt(54)

    h = doc.add_heading(f'Analise de Entradas ICMS/ES -- {periodo}', 0)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    inter = df[df['TIPO_OP'] == 'Interestadual']
    intra = df[df['TIPO_OP'] == 'Intraestadual']

    h1 = doc.add_heading('1. Visao Geral', 1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    tbl = doc.add_table(rows=7, cols=2)
    tbl.style = 'Table Grid'
    dados_vg = [
        ('Total de Registros', f'{len(df):,}'),
        ('Valor Total Entradas (R$)', f'{df["VLCONTABIL"].sum():,.2f}'),
        ('Base ICMS Total (R$)', f'{df["VLBASEICMS"].sum():,.2f}'),
        ('ICMS Total (R$)', f'{df["VLICMS"].sum():,.2f}'),
        ('Intraestadual (registros)', f'{len(intra):,}'),
        ('Interestadual (registros)', f'{len(inter):,}'),
        ('Periodo', periodo),
    ]
    for i, (lbl, val) in enumerate(dados_vg):
        tbl.rows[i].cells[0].text = lbl; tbl.rows[i].cells[1].text = str(val)
    doc.add_paragraph()

    infos = [
        ('DIV-1', 'Mercadoria Importada com Aliquota > 4%', 'CRITICA',
         'Mercadorias de origem estrangeira (Orig 1, 2 ou 6) tributadas com aliquota superior a 4% em operacoes interestaduais violam a Resolucao do Senado 13/2012. O credito escriturado a maior pode ser glosado pela SEFAZ-ES.',
         'Notificar os fornecedores para emissao de nota de correcao; estornar o credito indevido no EFD.'),
        ('DIV-2', 'CST 41 Interestadual -- Verificar Isencao', 'ATENCAO',
         'CST 41 em operacoes interestaduais indica isencao ou nao tributacao. Sem amparo em convenio CONFAZ especifico, o credito pode ser considerado indevido. A coexistencia de CST 41 e CST 00 no mesmo NCM/fornecedor e o cenario de maior risco.',
         'Solicitar ao fornecedor o convenio ou ato normativo que ampara a isencao; verificar se o mesmo produto aparece com CST 00 em outras NFs.'),
        ('DIV-3', 'CST 90 Interestadual -- Natureza Desconhecida', 'ATENCAO',
         'CST 90 em operacao interestadual indica tributacao "outras". O CFOP e a chave diagnostica: CFOP 2923 com valores altos representa complemento de preco ou bonificacao financeira.',
         'Mapear a natureza de cada operacao pelo CFOP; contatar fornecedor para esclarecimento; corrigir o CST conforme a natureza identificada.'),
        ('DIV-4', 'Base Reduzida sem CST 20', 'ATENCAO',
         'Base de calculo inferior ao valor do item com CST 00 declarado e inconsistente. O EFD exige CST 20 quando ha reducao de base por convenio CONFAZ. Confirmar o convenio aplicavel na SEFAZ-ES (www.sefaz.es.gov.br).',
         'Confirmar o convenio de amparo para cada cluster de reducao; corrigir a parametrizacao do ERP para emitir CST 20; retificar o EFD dos periodos afetados.'),
        ('DIV-5', 'Aliquota Atipica Interestadual', 'VERIFICAR',
         'Aliquotas fora do padrao interestadual (0%, 4%, 7% ou 12%) podem indicar erro de cadastro no ERP do fornecedor ou fornecedor local incorretamente classificado com CFOP 2xxx.',
         'Verificar o estado de origem do fornecedor; confirmar se ha convenio ou regime especial; corrigir o CFOP se o fornecedor for esense.'),
    ]

    for idx, (div_id, titulo, crit, narrativa, acao) in enumerate(infos, 2):
        chave = div_id.replace('-', '')
        dados = divs_dict.get(chave)
        n = len(dados) if dados is not None and not dados.empty else 0
        h2 = doc.add_heading(f'{div_id} -- {titulo}', 2)
        h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        p = doc.add_paragraph()
        clr = RGBColor(0xC0, 0x00, 0x00) if crit == 'CRITICA' else RGBColor(0xC5, 0x5A, 0x11) if crit == 'ATENCAO' else RGBColor(0x37, 0x56, 0x23)
        r2 = p.add_run(f'  {crit} -- {n} registros identificados')
        r2.bold = True; r2.font.color.rgb = clr
        doc.add_paragraph(narrativa)
        if dados is not None and not dados.empty and 'VLCONTABIL' in dados.columns:
            top5 = dados.nlargest(min(5, len(dados)), 'VLCONTABIL')
            doc.add_paragraph('Principais casos:')
            for _, row in top5.iterrows():
                nf = row.get('NUMNOTA', '--')
                vl = float(row.get('VLCONTABIL', row.get('VLITEM', 0)))
                desc = str(row.get('DESCRICAO', ''))[:50]
                doc.add_paragraph(f'NF {nf} -- R$ {vl:,.2f} -- {desc}', style='List Bullet')
        p2 = doc.add_paragraph()
        p2.add_run('Acao recomendada: ').bold = True
        p2.add_run(acao)
        doc.add_paragraph()

    h_res = doc.add_heading('Resumo e Prioridades', 1)
    h_res.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = 'Table Grid'
    for i, hdr in enumerate(['DIV', 'Criticidade', 'Registros', 'Dif. ICMS (R$)']):
        tbl2.rows[0].cells[i].text = hdr
    for div_id, _, crit, _, _ in infos:
        chave = div_id.replace('-', '')
        dados = divs_dict.get(chave)
        n = len(dados) if dados is not None and not dados.empty else 0
        dif = round(dados['Dif. ICMS (R$)'].sum(), 2) if dados is not None and not dados.empty and 'Dif. ICMS (R$)' in dados.columns else 0.0
        row_cells = tbl2.add_row().cells
        row_cells[0].text = div_id; row_cells[1].text = crit
        row_cells[2].text = str(n); row_cells[3].text = f'R$ {dif:,.2f}'

    doc.add_paragraph()
    p_rod = doc.add_paragraph(
        f'Gerado automaticamente em {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}'
        f' -- Especialista em Apuracao de ICMS'
    )
    if p_rod.runs:
        p_rod.runs[0].font.size = Pt(8)
        p_rod.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.save(str(caminho_word))


# -- main ----------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print('Uso: python analisar_entradas_es.py <arquivo1.xls> [arquivo2.csv] ...')
        sys.exit(1)

    caminhos = sys.argv[1:]
    print(f'Carregando {len(caminhos)} arquivo(s)...')
    df = carregar_dados(caminhos)
    print(f'  {len(df):,} registros | {df["CODFILIAL"].nunique()} filiais')

    periodo = extrair_periodo(df, caminhos)
    print(f'  Periodo: {periodo}')

    inter = df[df['TIPO_OP'] == 'Interestadual']

    print('Calculando divergencias...')
    d1 = calcular_div1(df, inter)
    d2 = calcular_div2(inter)
    d3 = calcular_div3(inter)
    d4 = calcular_div4(inter)
    d5 = calcular_div5(inter)

    divs = {'DIV1': d1, 'DIV2': d2, 'DIV3': d3, 'DIV4': d4, 'DIV5': d5}
    for k, v in divs.items():
        n = len(v) if v is not None and not v.empty else 0
        print(f'  {k}: {n} registros')

    caminho_excel = PASTA_SAIDA / f'Analise Entradas ICMS MT - {periodo}.xlsx'
    caminho_word  = PASTA_SAIDA / f'Analise Entradas ICMS MT - {periodo}.docx'

    print(f'Gerando Excel: {caminho_excel.name}')
    gerar_excel(df, divs, periodo, caminho_excel)

    print(f'Gerando Word: {caminho_word.name}')
    gerar_word(df, divs, periodo, caminho_word)

    print('Concluido!')


if __name__ == '__main__':
    main()
